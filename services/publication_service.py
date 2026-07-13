"""
services/publication_service.py
===============================
Service for exporting Tables (Markdown, LaTeX, CSV, DOCX),
formatting Citations, and building Reproducibility Packages.
"""

import csv
import io
import json
import sys
import platform
import uuid
from datetime import datetime
from typing import List, Dict, Any, Optional
from sqlalchemy.orm import Session
from database.models import (
    PublicationWorkspace, PublicationDraft, ReproducibilityLog,
    LiteratureReference, CampaignSeriesLink, SeriesCompound, ResearchProject,
    TableAsset, Portfolio, OptimizationCampaign
)

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import openpyxl
    from openpyxl.utils import get_column_letter
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False


class PublicationService:
    def __init__(self, db: Session):
        self.db = db

    # ===============================
    # Workspace & Drafts
    # ===============================
    
    def create_workspace(self, project_id: str, title: str, doc_type: str = "Manuscript") -> PublicationWorkspace:
        ws = PublicationWorkspace(
            id=f"pubws_{uuid.uuid4().hex[:10]}",
            project_id=project_id,
            title=title,
            type=doc_type
        )
        self.db.add(ws)
        self.db.commit()
        return ws

    def get_or_create_workspace(self, project_id: str, title: str = "Manuscript",
                                doc_type: str = "Manuscript") -> PublicationWorkspace:
        """One manuscript workspace per project by default; reused across drafts."""
        ws = (
            self.db.query(PublicationWorkspace)
            .filter_by(project_id=project_id, type=doc_type)
            .order_by(PublicationWorkspace.created_at.asc())
            .first()
        )
        if ws:
            return ws
        return self.create_workspace(project_id, title, doc_type)

    def project_id_for_campaign(self, campaign_id: str) -> Optional[str]:
        """Resolve Campaign -> Portfolio -> Project (the canonical ownership path)."""
        row = (
            self.db.query(Portfolio.project_id)
            .join(OptimizationCampaign, OptimizationCampaign.portfolio_id == Portfolio.id)
            .filter(OptimizationCampaign.id == campaign_id)
            .first()
        )
        return row[0] if row else None

    def save_draft(self, workspace_id: str, title: str, content_type: str, content: str, evidence: dict) -> PublicationDraft:
        """
        Persists a generated draft together with its evidence map. Version is
        auto-incremented per (workspace, content_type) so re-drafting keeps history.
        """
        prior = (
            self.db.query(PublicationDraft)
            .filter_by(workspace_id=workspace_id, content_type=content_type)
            .count()
        )
        draft = PublicationDraft(
            id=f"draft_{uuid.uuid4().hex[:10]}",
            workspace_id=workspace_id,
            title=title,
            content_type=content_type,
            content=content,
            evidence_links=evidence,
            version=prior + 1,
        )
        self.db.add(draft)
        self.db.commit()
        return draft

    def save_table_asset(self, workspace_id: str, title: str, table_type: str,
                         payload: Any, caption: str = "") -> TableAsset:
        """Persists an exported table so a manuscript can cite it later."""
        asset = TableAsset(
            id=f"tbl_{uuid.uuid4().hex[:10]}",
            workspace_id=workspace_id,
            title=title,
            table_type=table_type,
            data_payload=payload,
            caption=caption,
        )
        self.db.add(asset)
        self.db.commit()
        return asset

    def list_drafts(self, project_id: str) -> List[PublicationDraft]:
        ws_ids = [
            w.id for w in self.db.query(PublicationWorkspace).filter_by(project_id=project_id).all()
        ]
        if not ws_ids:
            return []
        return (
            self.db.query(PublicationDraft)
            .filter(PublicationDraft.workspace_id.in_(ws_ids))
            .order_by(PublicationDraft.created_at.desc())
            .all()
        )

    # ===============================
    # Citation Formatting
    # ===============================

    def format_citations(self, project_id: str, style: str = "ACS") -> List[Dict[str, str]]:
        """
        Formats all LiteratureReferences for a project into standard citation strings.
        Supports: ACS, APA, Nature.
        """
        refs = self.db.query(LiteratureReference).filter_by(project_id=project_id).order_by(LiteratureReference.year.asc()).all()
        formatted = []
        
        for i, ref in enumerate(refs):
            authors = ref.authors or "Unknown"
            year = ref.year or "n.d."
            title = ref.title or ""
            journal = ref.journal or ""
            
            if style == "ACS":
                # Author1; Author2. Title. Journal Year.
                citation = f"{authors}. {title}. <i>{journal}</i> <b>{year}</b>."
            elif style == "APA":
                # Author (Year). Title. Journal.
                citation = f"{authors} ({year}). {title}. <i>{journal}</i>."
            elif style == "Nature":
                # Author. Title. Journal Year (Nature uses numbered)
                citation = f"{i+1}. {authors}. {title}. <i>{journal}</i> ({year})."
            elif style == "Vancouver":
                # Author. Title. Journal. Year;
                citation = f"{i+1}. {authors}. {title}. {journal}. {year};"
            elif style == "IEEE":
                # [1] Author, "Title," Journal, Year.
                citation = f"[{i+1}] {authors}, \"{title},\" <i>{journal}</i>, {year}."
            else:
                citation = f"{authors} ({year}) {title}"
                
            formatted.append({
                "id": ref.id,
                "doi": ref.doi,
                "text": citation
            })
            
        return formatted

    # ===============================
    # Table Generators
    # ===============================

    def generate_compound_table(self, campaign_id: str, format_type: str = "csv") -> Any:
        """
        Generates a structured SAR/ADMET table for a given campaign.
        Formats: csv, latex, docx.
        """
        links = self.db.query(CampaignSeriesLink).filter_by(campaign_id=campaign_id).all()
        series_ids = [l.series_id for l in links]
        compounds = self.db.query(SeriesCompound).filter(SeriesCompound.series_id.in_(series_ids)).all()
        
        # Sort by pIC50 desc
        compounds.sort(key=lambda x: x.normalized_value or 0, reverse=True)
        
        headers = ["ID", "SMILES", "pIC50", "MW", "LogP", "TPSA", "Status"]
        rows = []
        for c in compounds:
            props = c.properties or {}
            rows.append([
                c.name,
                c.smiles,
                f"{c.normalized_value:.2f}" if c.normalized_value else "N/A",
                f"{props.get('mw', 0):.1f}",
                f"{props.get('logp', 0):.2f}",
                f"{props.get('tpsa', 0):.1f}",
                props.get("candidate_status", "Hit")
            ])
            
        if format_type == "csv":
            return self._to_csv(headers, rows)
        elif format_type == "latex":
            return self._to_latex(headers, rows)
        elif format_type == "docx":
            return self._to_docx(headers, rows)
        elif format_type == "excel":
            return self._to_excel(headers, rows)
        else:
            return None

    def _to_csv(self, headers: List[str], rows: List[List[str]]) -> str:
        output = io.StringIO()
        writer = csv.writer(output)
        writer.writerow(headers)
        writer.writerows(rows)
        return output.getvalue()
        
    def _to_latex(self, headers: List[str], rows: List[List[str]]) -> str:
        # Escape underscores for LaTeX
        def escape(val):
            return str(val).replace("_", "\\_").replace("%", "\\%")
            
        cols = "|".join(["c"] * len(headers))
        latex = f"\\begin{{table}}[htbp]\n\\centering\n\\begin{{tabular}}{{|{cols}|}}\n\\hline\n"
        latex += " & ".join([f"\\textbf{{{escape(h)}}}" for h in headers]) + " \\\\\n\\hline\n"
        
        for r in rows:
            latex += " & ".join([escape(val) for val in r]) + " \\\\\n"
            
        latex += "\\hline\n\\end{tabular}\n\\caption{Compound SAR Data}\n\\end{table}"
        return latex

    def _to_docx(self, headers: List[str], rows: List[List[str]]) -> Any:
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is not installed. Cannot generate DOCX file.")
            
        doc = Document()
        doc.add_heading('Compound Table', level=1)
        
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
            
        for r in rows:
            row_cells = table.add_row().cells
            for i, val in enumerate(r):
                row_cells[i].text = str(val)
                
        # Save to memory stream
        bio = io.BytesIO()
        doc.save(bio)
        return bio.getvalue()

    def _to_excel(self, headers: List[str], rows: List[List[str]]) -> Any:
        if not EXCEL_AVAILABLE:
            raise ImportError("openpyxl is not installed. Cannot generate Excel file.")
            
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Compound Data"
        
        ws.append(headers)
        for r in rows:
            ws.append(r)
            
        # Bold headers
        for cell in ws[1]:
            cell.font = openpyxl.styles.Font(bold=True)
            
        bio = io.BytesIO()
        wb.save(bio)
        return bio.getvalue()

    # ===============================
    # Reproducibility
    # ===============================

    def create_reproducibility_log(self, project_id: str, experiment_id: str, parameters: dict) -> ReproducibilityLog:
        """Captures environmental state + parameters."""
        env_snapshot = {
            "os": platform.platform(),
            "python": sys.version,
        }
        
        soft_versions = {}
        try:
            import rdkit
            soft_versions["rdkit"] = rdkit.__version__
        except ImportError:
            soft_versions["rdkit"] = "not_installed"
            
        # Normally would capture Vina version if executing Vina binary
        soft_versions["vina"] = "1.2.3 (assumed/bundled)"
        
        log = ReproducibilityLog(
            id=f"repr_{uuid.uuid4().hex[:10]}",
            project_id=project_id,
            experiment_id=experiment_id,
            software_versions=soft_versions,
            environment_snapshot=env_snapshot,
            parameters=parameters
        )
        self.db.add(log)
        self.db.commit()
        return log
